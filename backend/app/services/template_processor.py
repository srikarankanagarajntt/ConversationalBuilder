"""Template processor — fills Word templates with CV data using python-docx."""
from __future__ import annotations

import os
from typing import Any, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.cv_schema import CvSchema


class TemplateProcessor:
    """Processes Word templates and fills them with CV data."""

    TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "docs")
    
    # Template file names
    TEMPLATE_FILES = {
        "ntt-classic": "template.docx",
        "ntt-modern": "template.docx",
        "minimal": "template.docx",
    }

    def process_template(self, cv: CvSchema, template_id: str = "ntt-classic") -> Document:
        """
        Fill a Word template with CV data.
        
        Args:
            cv: CV schema data
            template_id: Template identifier
            
        Returns:
            Filled Document object
        """
        template_file = self.TEMPLATE_FILES.get(template_id, "template.docx")
        template_path = os.path.join(self.TEMPLATE_DIR, template_file)
        
        if not os.path.exists(template_path):
            # Fallback to creating a default document if template doesn't exist
            return self._create_default_document(cv, template_id)
        
        # Load template
        doc = Document(template_path)
        
        # Fill placeholders
        self._fill_placeholders(doc, cv)
        
        return doc

    def _fill_placeholders(self, doc: Document, cv: CvSchema) -> None:
        """Replace placeholder text in document with CV data."""
        
        # Create a mapping of placeholders to values
        replacements = {
            "{{fullName}}": cv.personalInfo.fullName or "Your Name",
            "{{email}}": cv.personalInfo.email or "",
            "{{phone}}": cv.personalInfo.phone or "",
            "{{location}}": cv.personalInfo.location or "",
            "{{linkedin}}": cv.personalInfo.linkedin or "",
            "{{summary}}": cv.personalInfo.summary or "",
            "{{skills}}": ", ".join(cv.skills) if cv.skills else "",
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    self._replace_text_in_paragraph(paragraph, key, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                self._replace_text_in_paragraph(paragraph, key, value)
        
        # Add detailed sections after template content
        self._add_experience_section(doc, cv)
        self._add_education_section(doc, cv)
        self._add_projects_section(doc, cv)
        self._add_certifications_section(doc, cv)
        self._add_languages_section(doc, cv)

    def _replace_text_in_paragraph(self, paragraph, key: str, value: str) -> None:
        """Replace placeholder text in a paragraph while preserving formatting."""
        if key in paragraph.text:
            # Clear paragraph
            paragraph.clear()
            # Add replaced text
            paragraph.add_run(value)

    def _add_experience_section(self, doc: Document, cv: CvSchema) -> None:
        """Add experience section to document."""
        if not cv.experience:
            return
        
        doc.add_heading("Professional Experience", level=1)
        
        for exp in cv.experience:
            # Job title and company
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{exp.title}")
            run.bold = True
            p.add_run(f" at {exp.company}").italic = True
            
            # Duration
            doc.add_paragraph(f"{exp.startDate} – {exp.endDate}", style='List Bullet 2')
            
            # Achievements
            for achievement in exp.achievements:
                doc.add_paragraph(achievement, style='List Bullet 2')

    def _add_education_section(self, doc: Document, cv: CvSchema) -> None:
        """Add education section to document."""
        if not cv.education:
            return
        
        doc.add_heading("Education", level=1)
        
        for edu in cv.education:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{edu.degree} in {edu.field}")
            run.bold = True
            p.add_run(f" from {edu.institution} ({edu.endDate})")

    def _add_projects_section(self, doc: Document, cv: CvSchema) -> None:
        """Add projects section to document."""
        if not cv.projects:
            return
        
        doc.add_heading("Projects", level=1)
        
        for project in cv.projects:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(project.name)
            run.bold = True
            
            if project.description:
                doc.add_paragraph(project.description, style='List Bullet 2')
            
            if project.technologies:
                tech_str = ", ".join(project.technologies)
                doc.add_paragraph(f"Technologies: {tech_str}", style='List Bullet 2')

    def _add_certifications_section(self, doc: Document, cv: CvSchema) -> None:
        """Add certifications section to document."""
        if not cv.certifications:
            return
        
        doc.add_heading("Certifications", level=1)
        
        for cert in cv.certifications:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(cert.name)
            run.bold = True
            p.add_run(f" from {cert.issuer} ({cert.date})")

    def _add_languages_section(self, doc: Document, cv: CvSchema) -> None:
        """Add languages section to document."""
        if not cv.languages:
            return
        
        doc.add_heading("Languages", level=1)
        doc.add_paragraph(", ".join(cv.languages))

    def _create_default_document(self, cv: CvSchema, template_id: str) -> Document:
        """Create a default CV document if template not found."""
        doc = Document()
        
        # Header with name
        heading = doc.add_heading(cv.personalInfo.fullName or "Curriculum Vitae", 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_lines = []
        if cv.personalInfo.email:
            contact_lines.append(cv.personalInfo.email)
        if cv.personalInfo.phone:
            contact_lines.append(cv.personalInfo.phone)
        if cv.personalInfo.location:
            contact_lines.append(cv.personalInfo.location)
        contact.add_run(" | ".join(contact_lines))
        
        # Summary
        if cv.personalInfo.summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(cv.personalInfo.summary)
        
        # Skills
        if cv.skills:
            doc.add_heading("Skills", level=1)
            skills_para = doc.add_paragraph()
            skills_para.add_run(", ".join(cv.skills))
        
        # Experience
        if cv.experience:
            doc.add_heading("Professional Experience", level=1)
            for exp in cv.experience:
                p = doc.add_paragraph(f"{exp.title} – {exp.company}", style='Heading 2')
                doc.add_paragraph(f"{exp.startDate} – {exp.endDate}", style='Heading 3')
                for achievement in exp.achievements:
                    doc.add_paragraph(achievement, style='List Bullet')
        
        # Education
        if cv.education:
            doc.add_heading("Education", level=1)
            for edu in cv.education:
                doc.add_paragraph(f"{edu.degree} in {edu.field} – {edu.institution} ({edu.endDate})", style='List Bullet')
        
        # Projects
        if cv.projects:
            doc.add_heading("Projects", level=1)
            for project in cv.projects:
                p = doc.add_paragraph(project.name, style='Heading 2')
                if project.description:
                    doc.add_paragraph(project.description)
                if project.technologies:
                    doc.add_paragraph(f"Technologies: {', '.join(project.technologies)}", style='List Bullet')
        
        # Certifications
        if cv.certifications:
            doc.add_heading("Certifications", level=1)
            for cert in cv.certifications:
                doc.add_paragraph(f"{cert.name} – {cert.issuer} ({cert.date})", style='List Bullet')
        
        # Languages
        if cv.languages:
            doc.add_heading("Languages", level=1)
            doc.add_paragraph(", ".join(cv.languages))
        
        return doc
